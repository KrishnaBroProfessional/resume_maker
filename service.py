from model import UserProfileDB
from ai_service import generate_resume_fields
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import markdown2
# from weasyprint import HTML
import io
# import os
# from dotenv import load_dotenv

# load_dotenv()

# Define MongoDB Connection String
MONGO_CONNECTION_STRING = "mongodb://localhost:27017"

# Initialize Database Connection
db = UserProfileDB(MONGO_CONNECTION_STRING)

# Service functions mapped to model.py functions

def create_profile(profile_data):
    return db.create_profile(profile_data)

def update_field(name, field, value):
    return db.update_field(name, field, value)

def get_field(name, field):
    return db.get_field(name, field)

def set_bio(name, bio):
    return db.set_bio(name, bio)

def get_bio(name):
    return db.get_bio(name)

def set_profile_pic(name, profile_pic):
    return db.set_profile_pic(name, profile_pic)

def get_profile_pic(name):
    return db.get_profile_pic(name)

def set_objective(name, objective):
    return db.set_objective(name, objective)

def get_objective(name):
    return db.get_objective(name)

def set_experience(name, experience):
    return db.set_experience(name, experience)

def get_experience(name):
    return db.get_experience(name)

def set_projects(name, projects):
    return db.set_projects(name, projects)

def get_projects(name):
    return db.get_projects(name)

def set_skills(name, skills):
    return db.set_skills(name, skills)

def get_skills(name):
    return db.get_skills(name)

def set_certifications(name, certifications):
    return db.set_certifications(name, certifications)

def get_certifications(name):
    return db.get_certifications(name)

def set_internships(name, internships):
    return db.set_internships(name, internships)

def get_internships(name):
    return db.get_internships(name)

def set_education(name, education):
    return db.set_education(name, education)

def get_education(name):
    return db.get_education(name)

def set_achievements(name, achievements):
    return db.set_achievements(name, achievements)

def get_achievements(name):
    return db.get_achievements(name)

def set_research_work(name, research_work):
    return db.set_research_work(name, research_work)

def get_research_work(name):
    return db.get_research_work(name)

# Additional function as requested
def get_job_post(job_str, name):
    """
    Function to get job post based on input string and name.
    Implementation can be customized as needed.
    """
    # Placeholder implementation: return a dict with inputs
    return {
        "job_str": job_str,
        "name": name,
        "message": "This is a placeholder response for get_job_post."
    }

def generate_resume_for_profile(job_post: str, profile_data: dict) -> dict:
    """
    Generate resume fields for a given job post and profile data using AI service.
    """
    return generate_resume_fields(job_post, profile_data)

def generate_docx_resume(resume_markdown: str) -> bytes:
    """
    Generate a DOCX file bytes from markdown resume text with basic formatting.
    """
    # Convert markdown to HTML
    html = markdown2.markdown(resume_markdown)

    # Create a new Document
    doc = Document()

    # Simple HTML to DOCX conversion for paragraphs and basic formatting
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")

    for element in soup.children:
        if element.name == 'h1':
            p = doc.add_heading(level=1)
            p.add_run(element.get_text())
        elif element.name == 'h2':
            p = doc.add_heading(level=2)
            p.add_run(element.get_text())
        elif element.name == 'h3':
            p = doc.add_heading(level=3)
            p.add_run(element.get_text())
        elif element.name == 'p':
            p = doc.add_paragraph()
            p.add_run(element.get_text())
        elif element.name == 'ul':
            for li in element.find_all('li'):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(li.get_text())
        elif element.name == 'ol':
            for li in element.find_all('li'):
                p = doc.add_paragraph(style='List Number')
                p.add_run(li.get_text())
        else:
            # For other tags, add plain text
            doc.add_paragraph(element.get_text())

    buffer = BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    return docx_bytes

# def generate_pdf_resume(resume_markdown: str) -> bytes:
#     """
#     Generate a PDF file bytes from markdown resume text with basic formatting.
#     Uses markdown2 to convert markdown to HTML and WeasyPrint to convert HTML to PDF.
#     """
#     # Convert markdown to HTML
#     html = markdown2.markdown(resume_markdown)

#     # Use WeasyPrint to convert HTML to PDF in memory
#     pdf_io = io.BytesIO()
#     HTML(string=html).write_pdf(pdf_io)
#     pdf_bytes = pdf_io.getvalue()
#     pdf_io.close()

#     return pdf_bytes

def generate_pdf_resume(resume_text: str) -> bytes:
    """
    Generate a PDF file bytes from the resume text.
    """
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    text_object = c.beginText(40, height - 40)
    lines = resume_text.split('\n')
    for line in lines:
        text_object.textLine(line)
    c.drawText(text_object)
    c.showPage()
    c.save()
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes
